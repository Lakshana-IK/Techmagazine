# Updated report.py to generate only the summary and prediction details in DOCX
from docx import Document
import os

def create_report(summary_text, chart_paths, output_path, table_df=None):
    doc = Document()
    doc.add_heading("🌾 Agricultural Market Analysis Report", level=1)

    doc.add_heading("Summary", level=2)
    for line in summary_text.strip().splitlines():
        doc.add_paragraph(line.strip())

    doc.save(output_path)